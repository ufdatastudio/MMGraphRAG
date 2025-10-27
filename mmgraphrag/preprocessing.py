from typing import Callable, Dict, List, Optional, Type, Union, cast
from dataclasses import dataclass
from docx import Document
from PIL import Image
from io import BytesIO
from parameter import cache_path
import asyncio
import os
import json
import base64

from base import (
    logger,
    compute_mdhash_id,
    encode_string_by_tiktoken,
    decode_tokens_by_tiktoken,
    read_config_to_dict,
)

from storage import (
    BaseKVStorage,
    JsonKVStorage,
    StorageNameSpace,
)
from llm import multimodel_if_cache
from prompt import PROMPTS

global_config_path = os.path.join(cache_path, "global_config.csv")

def chunking_by_token_size(
    content: str, overlap_token_size=128, max_token_size=1024, tiktoken_model="gpt-4o"
):
    """
    Chunk text by token size.

    This function is used to chunk given text content according to specified token size limits,
    while ensuring overlap between adjacent chunks. Mainly used for processing large texts to fit
    input limits of models like OpenAI's GPT series.

    Parameters:
    - content: str, text content to be chunked.
    - overlap_token_size: int, default 128. Number of overlapping tokens between adjacent text chunks.
    - max_token_size: int, default 1024. Maximum number of tokens per text chunk.
    - tiktoken_model: str, default "gpt-4o". Tiktoken model used for tokenization and detokenization.

    Returns:
    - List[Dict[str, Any]], list containing number of tokens, text content, and chunk order index for each chunk.
    """
    # Tokenize text using specified tiktoken model
    tokens = encode_string_by_tiktoken(content, model_name=tiktoken_model)
    # Initialize list to store chunking results
    results = []
    # Traverse tokens and chunk according to max_token_size and overlap_token_size
    for index, start in enumerate(
        range(0, len(tokens), max_token_size - overlap_token_size)
    ):
        # Get chunk tokens based on current chunk start position and max token limit
        chunk_content = decode_tokens_by_tiktoken(
            tokens[start : start + max_token_size], model_name=tiktoken_model
        )
        # Add current chunk's token count, text content, and chunk order index to results list
        results.append(
            {
                "tokens": min(max_token_size, len(tokens) - start),
                "content": chunk_content.strip(),
                "chunk_order_index": index,
            }
        )
    return results

@dataclass
class text_chunking_func:
    # Chunking
    chunk_func: Callable[[str, Optional[int], Optional[int], Optional[str]], List[Dict[str, Union[str, int]]]] = chunking_by_token_size
    # Chunk size
    chunk_token_size: int = 1200
    # Chunk overlap quantity
    chunk_overlap_token_size: int = 100
    

    # Key-value storage, json, specifically defined in storage.py
    key_string_value_json_storage_cls: Type[BaseKVStorage] = JsonKVStorage
    # Get global settings
    global_config = read_config_to_dict(global_config_path)
    
    # tiktoken model name, default gpt-4o, moonshot-v1-32k is compatible
    tiktoken_model_name = global_config["tiktoken_model_name"]

    def __post_init__(self):
        # Initialize storage class instance for storing full documents
        self.full_docs = self.key_string_value_json_storage_cls(
            namespace="full_docs", global_config = self.global_config
        )
        # Initialize storage class instance for storing text chunks
        self.text_chunks = self.key_string_value_json_storage_cls(
            namespace="text_chunks", global_config = self.global_config
        )
    
    async def text_chunking(self,string_or_strings):
        try:
            # If input is a string, convert it to a list
            if isinstance(string_or_strings, str):
                string_or_strings = [string_or_strings]
            # ---------- new docs
            # For each element in string_or_strings, strip whitespace and use as document content.
            # Calculate its MD5 hash with "doc-" prefix as key, content itself as value, generating a new dict new_docs.
            new_docs = {
                compute_mdhash_id(c.strip(), prefix="doc-"): {"content": c.strip()}
                for c in string_or_strings
            }
            # Filter out new document IDs that need to be added
            _add_doc_keys = await self.full_docs.filter_keys(list(new_docs.keys()))
            # Update new documents dict based on filtering results
            new_docs = {k: v for k, v in new_docs.items() if k in _add_doc_keys}
            # If no new documents need to be added, log and return
            if not len(new_docs):
                logger.warning(f"All docs are already in the storage")
                return
            # Log insertion of new documents
            logger.info(f"[New Docs] inserting {len(new_docs)} docs")

            # ---------- chunking
            inserting_chunks = {}
            for doc_key, doc in new_docs.items():
                # Generate chunks for each document
                chunks = {
                    compute_mdhash_id(dp["content"], prefix="chunk-"): {
                        **dp,
                        "full_doc_id": doc_key,
                    }
                    for dp in self.chunk_func(
                        doc["content"],
                        overlap_token_size=self.chunk_overlap_token_size,
                        max_token_size=self.chunk_token_size,
                        tiktoken_model=self.tiktoken_model_name,
                    )
                }
                inserting_chunks.update(chunks)
            # Filter out new chunk IDs that need to be added
            _add_chunk_keys = await self.text_chunks.filter_keys(
                list(inserting_chunks.keys())
            )
            # Update new chunks dict based on filtering results
            inserting_chunks = {
                k: v for k, v in inserting_chunks.items() if k in _add_chunk_keys
            }
            # If no new chunks need to be added, log and return
            if not len(inserting_chunks):
                logger.warning(f"All chunks are already in the storage")
                return
            logger.info(f"[New Chunks] inserting {len(inserting_chunks)} chunks")

            # Submit all updates and index operations
            await self.full_docs.upsert(new_docs)
            await self.text_chunks.upsert(inserting_chunks)
        finally:
            await self._text_chunking_done()
    async def _text_chunking_done(self):
        tasks = []
        for storage_inst in [
            self.full_docs,
            self.text_chunks,
        ]:
            if storage_inst is None:
                continue
            tasks.append(cast(StorageNameSpace, storage_inst).index_done_callback())
        await asyncio.gather(*tasks)

image_description_prompt_user = PROMPTS["image_description_user"]
image_description_prompt_system = PROMPTS["image_description_system"]

async def get_image_description(image_path):
    with open(image_path, 'rb') as img_file:
        image_base = base64.b64encode(img_file.read()).decode('utf-8')
    image_description =  await multimodel_if_cache(user_prompt=image_description_prompt_user,img_base=image_base,system_prompt=image_description_prompt_system)
    return image_description

def find_chunk_for_image(text_chunks, preceding_text, following_text):
    """
    Find the chunk an image belongs to based on text before and after the image.
    Prioritize chunks containing more continuous characters, ignoring newlines.
    """
    best_chunk_id = None
    best_match_count = 0

    # Combine preceding and following text as a whole, remove newlines
    combined_text = f"{preceding_text} {following_text}".replace('\n', '').strip()

    # If combined text is empty, return None
    if not combined_text:
        return None

    # Traverse all chunks
    for chunk_id, chunk_data in text_chunks.items():
        # Remove newlines from chunk
        chunk_content = chunk_data['content'].replace('\n', '')

        # Calculate match score between combined text and chunk content (based on word matching)
        match_count = sum(1 for word in combined_text.split() if word in chunk_content)

        # If current chunk has highest match score, select it
        if match_count > best_match_count:
            best_match_count = match_count
            best_chunk_id = chunk_id

    return best_chunk_id

def extract_image_context(docx_path, context_length=100):
    """
    Extract text information before and after images in a document.
    """
    doc = Document(docx_path)
    results = {}
    image_counter = 1  # Initialize image counter

    def get_previous_text(index, length_needed):
        """
        Recursively get text from preceding paragraphs until length_needed characters are reached.
        """
        text = ""
        while index >= 0 and len(text) < length_needed:
            paragraph_text = doc.paragraphs[index].text
            text = paragraph_text[-(length_needed - len(text)):] + text
            index -= 1
        return text

    def get_next_text(index, length_needed):
        """
        Recursively get text from following paragraphs until length_needed characters are reached.
        """
        text = ""
        while index < len(doc.paragraphs) and len(text) < length_needed:
            paragraph_text = doc.paragraphs[index].text
            text = text + paragraph_text[:length_needed - len(text)]
            index += 1
        return text

    for i, paragraph in enumerate(doc.paragraphs):
        # Get text from each paragraph
        text = paragraph.text

        # Check if paragraph contains images
        has_image = any(run._element.xpath('.//a:blip') for run in paragraph.runs)

        if has_image:
            # Extract text before and after image

            # Get text from preceding paragraph
            before_text = get_previous_text(i - 1, context_length)

            # Get text from following paragraph
            after_text = get_next_text(i + 1, context_length)

            # Save results, ensure before and after text are not empty
            results[f"image_{image_counter}"] = {
                "before": before_text.strip(),
                "after": after_text.strip()
            }

            # Increment image counter
            image_counter += 1

    return results

def compress_image_to_size(input_image, output_image_path, target_size_mb=5, step=10, quality=90):
    """
    Compress image to within target size (in MB).

    Parameters:
    input_image (PIL.Image): Input image PIL object.
    output_image_path (str): Output image path.
    target_size_mb (int): Target size in MB, default 5MB.
    step (int): Quality reduction step size per iteration, default 10.
    quality (int): Initial saved image quality, default 90.

    Returns:
    bool: Whether successfully compressed to within target size.
    """
    target_size_bytes = target_size_mb * 1024 * 1024  # Convert target size to bytes

    # First save image and check size
    img = input_image
    img.save(output_image_path, quality=quality)
    if os.path.getsize(output_image_path) <= target_size_bytes:
        return True

    # Try gradually reducing quality until image size is below target
    while os.path.getsize(output_image_path) > target_size_bytes and quality > 10:
        quality -= step
        img.save(output_image_path, quality=quality)
    
    # Check if final size meets requirements
    if os.path.getsize(output_image_path) <= target_size_bytes:
        return True
    else:
        logger.warning("Unable to compress image to within target size, please adjust initial quality or step size in preprocessing.py.")
        return False

async def extract_text_and_images_with_chunks(docx_path, output_dir,context_length):
    """
    Extract text chunks from document and associate with images. Image context text extraction and integration.
    """
    doc = Document(docx_path)

    # Extract all text content from document
    full_text = "\n".join([para.text for para in doc.paragraphs])

    # First instantiate class
    text_chunking_instance = text_chunking_func()
    # Chunk document text
    await text_chunking_instance.text_chunking(full_text)
    filepath = os.path.join(output_dir, 'kv_store_text_chunks.json')
    with open(filepath, 'r') as file:
        text_chunks = json.load(file)

    # Create directory to save images
    images_dir = os.path.join(output_dir, 'images')
    if not os.path.exists(images_dir):
        os.makedirs(images_dir)

    # Extract image context (text before and after)
    image_contexts = extract_image_context(docx_path, context_length)

    # Generate final result dictionary
    image_data = {}
    image_count = 0
    extracted_images = []  # Track extracted images

    # Step 1: Extract images using XML parsing (in order)
    for idx, shape in enumerate(doc.element.xpath("//w:drawing//a:blip"), start=1):
        # Extract `r:embed` attribute which points to the image relationship ID
        embed = shape.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if embed in doc.part.rels:
            image_count += 1
            image = doc.part.rels[embed].target_part.blob

            # Use PIL to open image and convert to JPG format
            image_bytes = BytesIO(image)
            with Image.open(image_bytes) as img:
                rgb_img = img.convert('RGB')  # Convert to RGB mode, suitable for JPG format
                image_output_path = os.path.join(images_dir, f'image_{image_count}.jpg')
                # Use compression function to save image smaller than target_size_mb
                compress_image_to_size(rgb_img, image_output_path, target_size_mb=5)

            # Record image information for subsequent processing
            extracted_images.append({
                'image_id': image_count,
                'image_path': image_output_path,
                'image_description': await get_image_description(image_output_path)
            })

    # Step 2: Associate image context
    for image_key, context in image_contexts.items():
        # Extract image index from image_key (avoid index out of range)
        image_index = int(image_key.split('_')[1]) - 1
        if image_index < len(extracted_images):
            current_image = extracted_images[image_index]  # Get current image
            preceding_text_segment = context["before"]
            following_text_segment = context["after"]
            
            
            # Find chunk the image belongs to
            associated_chunk_id = find_chunk_for_image(text_chunks, preceding_text_segment, following_text_segment)
            if associated_chunk_id:
                # Add image information to image_data
                image_data[image_key] = {
                    "chunk_order_index": text_chunks[associated_chunk_id]['chunk_order_index'],
                    "chunk_id": associated_chunk_id,
                    "image_id": current_image['image_id'],
                    "description": current_image['image_description'],
                    "image_path": current_image['image_path'],
                    "before": preceding_text_segment,
                    "after": following_text_segment
                }
    return image_data

@dataclass
class chunking_func:
    # Image extraction context length (100 each, so total length is 200)
    context_length: int = 100

    # Key-value storage, json, specifically defined in storage.py
    key_string_value_json_storage_cls: Type[BaseKVStorage] = JsonKVStorage

    # Get global settings
    global_config = read_config_to_dict(global_config_path)

    def __post_init__(self):
        # Initialize storage class instance for storing image attributes
        self.image_data = self.key_string_value_json_storage_cls(
            namespace="image_data", global_config = self.global_config
        )
    
    async def extract_text_and_images(self,docx_path):
        try:
            output_dir = self.global_config["working_dir"]
            context_length = self.context_length
            imagedata = await extract_text_and_images_with_chunks(docx_path, output_dir, context_length)
            # Submit all updates and index operations
            await self.image_data.upsert(imagedata)
        finally:
            await self._chunking_done()
    async def _chunking_done(self):
        tasks = []
        for storage_inst in [
            self.image_data
        ]:
            if storage_inst is None:
                continue
            tasks.append(cast(StorageNameSpace, storage_inst).index_done_callback())
        await asyncio.gather(*tasks)